import re
from pathlib import Path

from docx import Document
from docx.opc.exceptions import PackageNotFoundError


PADRAO_VARIAVEL = re.compile(
    r"\{\{\s*([^{}]+?)\s*\}\}"
)


class ErroLeituraDocumento(Exception):
    """Erro ao abrir ou analisar um arquivo DOCX."""


def extrair_variaveis_texto(texto):
    """
    Encontra variáveis no formato {{variavel}} em um texto.

    Exemplo:
        Olá, {{ cliente.nome }}.

    Resultado:
        ["cliente.nome"]
    """
    if not texto:
        return []

    variaveis = []

    for resultado in PADRAO_VARIAVEL.findall(texto):
        variavel = resultado.strip()

        if variavel:
            variaveis.append(variavel)

    return variaveis


def extrair_texto_paragrafos(paragrafos):
    """
    Junta o conteúdo de uma coleção de parágrafos.

    A união é importante porque o Word pode dividir uma variável
    em vários trechos internos, chamados runs.
    """
    textos = []

    for paragrafo in paragrafos:
        texto = "".join(
            run.text
            for run in paragrafo.runs
        )

        if not texto:
            texto = paragrafo.text or ""

        textos.append(texto)

    return "\n".join(textos)


def extrair_texto_tabela(tabela):
    """
    Extrai o conteúdo de todas as células de uma tabela,
    incluindo tabelas aninhadas.
    """
    textos = []

    for linha in tabela.rows:
        for celula in linha.cells:
            textos.append(
                extrair_texto_paragrafos(
                    celula.paragraphs
                )
            )

            for tabela_interna in celula.tables:
                textos.append(
                    extrair_texto_tabela(
                        tabela_interna
                    )
                )

    return "\n".join(textos)


def extrair_texto_secao(secao):
    """
    Extrai cabeçalho e rodapé de uma seção do documento.
    """
    textos = []

    cabecalhos = [
        secao.header,
        secao.first_page_header,
        secao.even_page_header,
    ]

    rodapes = [
        secao.footer,
        secao.first_page_footer,
        secao.even_page_footer,
    ]

    for area in cabecalhos + rodapes:
        textos.append(
            extrair_texto_paragrafos(
                area.paragraphs
            )
        )

        for tabela in area.tables:
            textos.append(
                extrair_texto_tabela(
                    tabela
                )
            )

    return "\n".join(textos)


def extrair_texto_documento(documento):
    """
    Extrai o texto que pode conter variáveis no documento inteiro.
    """
    textos = [
        extrair_texto_paragrafos(
            documento.paragraphs
        )
    ]

    for tabela in documento.tables:
        textos.append(
            extrair_texto_tabela(
                tabela
            )
        )

    for secao in documento.sections:
        textos.append(
            extrair_texto_secao(
                secao
            )
        )

    return "\n".join(textos)


def escanear_variaveis_docx(caminho_arquivo):
    """
    Abre um DOCX e retorna uma lista ordenada e sem duplicações.

    Exemplo:
        [
            "caso.numero_interno",
            "cliente.cpf",
            "cliente.nome",
        ]
    """
    caminho = Path(caminho_arquivo)

    if not caminho.exists():
        raise ErroLeituraDocumento(
            "O arquivo informado não foi encontrado."
        )

    if not caminho.is_file():
        raise ErroLeituraDocumento(
            "O caminho informado não corresponde a um arquivo."
        )

    if caminho.suffix.lower() != ".docx":
        raise ErroLeituraDocumento(
            "O scanner aceita somente arquivos DOCX."
        )

    try:
        documento = Document(str(caminho))

    except PackageNotFoundError as erro:
        raise ErroLeituraDocumento(
            "O arquivo DOCX está inválido ou corrompido."
        ) from erro

    except Exception as erro:
        raise ErroLeituraDocumento(
            "Não foi possível analisar o arquivo DOCX."
        ) from erro

    texto_completo = extrair_texto_documento(
        documento
    )

    variaveis = extrair_variaveis_texto(
        texto_completo
    )

    return sorted(
        set(variaveis),
        key=lambda item: item.lower(),
    )