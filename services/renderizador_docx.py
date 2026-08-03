import re
from copy import deepcopy
from pathlib import Path
from typing import Any

from docx import Document
from docx.document import Document as DocumentoDocx
from docx.table import _Cell, Table
from docx.text.paragraph import Paragraph


PADRAO_VARIAVEL = re.compile(
    r"\{\{\s*([^{}]+?)\s*\}\}"
)


class ErroRenderizacaoDocumento(Exception):
    """
    Erro ocorrido durante a geração do documento DOCX.
    """


def normalizar_contexto(
    contexto: dict[str, Any],
) -> dict[str, str]:
    """
    Garante que todas as chaves e valores estejam
    em formato textual.
    """
    contexto_normalizado: dict[str, str] = {}

    for chave, valor in contexto.items():
        codigo = str(chave).strip()

        if not codigo:
            continue

        if valor is None:
            texto = ""
        else:
            texto = str(valor)

        contexto_normalizado[codigo] = texto

    return contexto_normalizado


def substituir_variaveis_texto(
    texto: str,
    contexto: dict[str, str],
    manter_nao_encontradas: bool = True,
) -> tuple[str, set[str]]:
    """
    Substitui as variáveis encontradas em um texto.

    Retorna:
    - o texto já substituído;
    - as variáveis que não foram encontradas no contexto.
    """
    nao_encontradas: set[str] = set()

    def substituir(
        correspondencia: re.Match,
    ) -> str:
        codigo = (
            correspondencia.group(1)
            .strip()
        )

        if codigo in contexto:
            return contexto[codigo]

        nao_encontradas.add(codigo)

        if manter_nao_encontradas:
            return correspondencia.group(0)

        return ""

    novo_texto = PADRAO_VARIAVEL.sub(
        substituir,
        texto,
    )

    return novo_texto, nao_encontradas


def substituir_em_paragrafo(
    paragrafo: Paragraph,
    contexto: dict[str, str],
    manter_nao_encontradas: bool = True,
) -> set[str]:
    """
    Substitui variáveis em um parágrafo.

    A função também consegue substituir variáveis que o Word
    dividiu em vários trechos internos, chamados runs.
    """
    nao_encontradas: set[str] = set()

    if not paragrafo.runs:
        return nao_encontradas

    texto_original = "".join(
        run.text
        for run in paragrafo.runs
    )

    if "{{" not in texto_original:
        return nao_encontradas

    texto_substituido, faltantes = (
        substituir_variaveis_texto(
            texto_original,
            contexto,
            manter_nao_encontradas,
        )
    )

    nao_encontradas.update(faltantes)

    if texto_substituido == texto_original:
        return nao_encontradas

    primeiro_run = paragrafo.runs[0]

    primeiro_run.text = texto_substituido

    for run in paragrafo.runs[1:]:
        run.text = ""

    return nao_encontradas


def substituir_em_tabela(
    tabela: Table,
    contexto: dict[str, str],
    manter_nao_encontradas: bool = True,
) -> set[str]:
    """
    Substitui variáveis em todas as células de uma tabela.
    """
    nao_encontradas: set[str] = set()

    for linha in tabela.rows:
        for celula in linha.cells:
            nao_encontradas.update(
                substituir_em_container(
                    celula,
                    contexto,
                    manter_nao_encontradas,
                )
            )

    return nao_encontradas


def substituir_em_container(
    container: DocumentoDocx | _Cell,
    contexto: dict[str, str],
    manter_nao_encontradas: bool = True,
) -> set[str]:
    """
    Substitui variáveis nos parágrafos e tabelas
    de um documento ou célula.
    """
    nao_encontradas: set[str] = set()

    for paragrafo in container.paragraphs:
        nao_encontradas.update(
            substituir_em_paragrafo(
                paragrafo,
                contexto,
                manter_nao_encontradas,
            )
        )

    for tabela in container.tables:
        nao_encontradas.update(
            substituir_em_tabela(
                tabela,
                contexto,
                manter_nao_encontradas,
            )
        )

    return nao_encontradas


def substituir_em_cabecalhos_rodapes(
    documento: DocumentoDocx,
    contexto: dict[str, str],
    manter_nao_encontradas: bool = True,
) -> set[str]:
    """
    Substitui variáveis nos cabeçalhos e rodapés
    de todas as seções.
    """
    nao_encontradas: set[str] = set()

    partes_processadas: set[int] = set()

    for secao in documento.sections:
        partes = [
            secao.header,
            secao.first_page_header,
            secao.even_page_header,
            secao.footer,
            secao.first_page_footer,
            secao.even_page_footer,
        ]

        for parte in partes:
            identificador = id(
                parte._element
            )

            if identificador in partes_processadas:
                continue

            partes_processadas.add(
                identificador
            )

            nao_encontradas.update(
                substituir_em_container(
                    parte,
                    contexto,
                    manter_nao_encontradas,
                )
            )

    return nao_encontradas


def renderizar_documento_docx(
    caminho_modelo: str | Path,
    caminho_saida: str | Path,
    contexto: dict[str, Any],
    manter_nao_encontradas: bool = True,
) -> dict[str, Any]:
    """
    Abre um modelo DOCX, substitui as variáveis
    e salva um novo arquivo.

    Retorna informações sobre o arquivo gerado.
    """
    caminho_modelo = Path(
        caminho_modelo
    )

    caminho_saida = Path(
        caminho_saida
    )

    if not caminho_modelo.exists():
        raise ErroRenderizacaoDocumento(
            "O arquivo do modelo não foi encontrado."
        )

    if not caminho_modelo.is_file():
        raise ErroRenderizacaoDocumento(
            "O caminho informado não é um arquivo válido."
        )

    if (
        caminho_modelo.suffix.lower()
        != ".docx"
    ):
        raise ErroRenderizacaoDocumento(
            "O modelo precisa estar no formato DOCX."
        )

    contexto_normalizado = normalizar_contexto(
        contexto
    )

    try:
        documento = Document(
            str(caminho_modelo)
        )

    except Exception as erro:
        raise ErroRenderizacaoDocumento(
            "Não foi possível abrir o arquivo DOCX."
        ) from erro

    nao_encontradas: set[str] = set()

    nao_encontradas.update(
        substituir_em_container(
            documento,
            contexto_normalizado,
            manter_nao_encontradas,
        )
    )

    nao_encontradas.update(
        substituir_em_cabecalhos_rodapes(
            documento,
            contexto_normalizado,
            manter_nao_encontradas,
        )
    )

    try:
        caminho_saida.parent.mkdir(
            parents=True,
            exist_ok=True,
        )

        documento.save(
            str(caminho_saida)
        )

    except Exception as erro:
        raise ErroRenderizacaoDocumento(
            "Não foi possível salvar o documento gerado."
        ) from erro

    return {
        "caminho": str(caminho_saida),
        "nome_arquivo": caminho_saida.name,
        "variaveis_nao_encontradas": sorted(
            nao_encontradas
        ),
        "quantidade_nao_encontradas": len(
            nao_encontradas
        ),
        "sucesso": True,
    }